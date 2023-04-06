from docx import Document
from docx.shared import Inches
from PIL import Image
import io
from tqdm import tqdm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.shared import Pt


class WordSaver:
    def __int__(self):
        pass

    def inches_from_pixels(self, pixels, dpi=150):
        return pixels / dpi

    def save(self, content_list):
        # 创建一个新的Word文档
        doc = Document()

        style_name = 'CustomStyle'
        custom_style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

        # 设置行间距
        custom_style.paragraph_format.space_before = Pt(0.8)  # 设置段落前的行间距
        custom_style.paragraph_format.space_after = Pt(0.8)  # 设置段落后的行间距

        # 遍历列表，根据元素类型添加到文档中
        for item in tqdm(content_list):
            if isinstance(item, str):
                # 如果元素是字符串，添加一个新的段落
                if item == "TRY_TO_BREAK":

                    paragraph = doc.add_paragraph()
                    paragraph.style = 'Normal'
                    paragraph.add_run('This paragraph is on a new page.')
                    paragraph.runs[0].add_break(WD_BREAK.PAGE)
                else:
                    doc.add_paragraph(item, style=style_name)

            elif isinstance(item, Image.Image):
                # 如果元素是PIL.Image对象，先将其保存到内存缓冲区，再将其添加到文档中
                with io.BytesIO() as buf:
                    item.save(buf, format='PNG')
                    buf.seek(0)
                    doc.add_picture(buf, width=Inches(self.inches_from_pixels(item.width)),
                                    height=Inches(self.inches_from_pixels(item.height)))  # 设置图片宽度，如有需要可以调整

        # 保存Word
        doc.save('demo.docx')
